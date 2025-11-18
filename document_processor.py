import os
import pypdf  # Changed from PyPDF2
from docx import Document
from typing import List, Dict, Any

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.pptx'}
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process a single file and extract text content"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._process_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._process_docx(file_path)
            elif file_ext == '.txt':
                return self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            return {"error": f"Error processing {file_path}: {str(e)}", "content": ""}
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)  # Changed from PyPDF2
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        return {
            "content": text,
            "pages": len(pdf_reader.pages),
            "type": "pdf"
        }
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return {
            "content": text,
            "paragraphs": len(doc.paragraphs),
            "type": "docx"
        }
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        return {
            "content": text,
            "type": "txt"
        }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - chunk_overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            if i + chunk_size >= len(words):
                break
        
        return chunks